from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports/1c_vat_legislation_audit.docx")


COLORS = {
    "blue": "1F4E79",
    "blue2": "2E75B6",
    "dark": "1F2937",
    "muted": "667085",
    "light_blue": "D9EAF7",
    "pale_blue": "EEF6FC",
    "pale_gray": "F4F6F8",
    "green": "E2F0D9",
    "yellow": "FFF2CC",
    "red": "FCE4D6",
    "border": "B7C9D6",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border"], size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, bold=False, color=None, size=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    style_run(run, bold=True, color=COLORS["blue"], size=24)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    style_run(run, color=COLORS["muted"], size=12)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        style_run(run, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc: Document, label: str, text: str, fill: str = COLORS["pale_blue"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    style_run(r, bold=True, color=COLORS["blue"])
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              risk_column: int | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(header_cells[i], COLORS["light_blue"])
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        style_run(r, bold=True, color=COLORS["dark"], size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if risk_column is not None and i == risk_column:
                fill = COLORS["pale_gray"]
                if "Высокий" in text:
                    fill = COLORS["red"]
                elif "Средний" in text:
                    fill = COLORS["yellow"]
                elif "Низкий" in text or "ОК" in text:
                    fill = COLORS["green"]
                set_cell_shading(cells[i], fill)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            style_run(r, size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue2"], 12, 6),
        ("Heading 3", 12, COLORS["blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("1C AI Auditor - VAT static compliance review")
    return doc


def build() -> None:
    doc = setup_doc()
    add_title(
        doc,
        "Статический аудит конфигурации 1С по НДС",
        "Сверка метаданных и BSL через MCP 1c-meta с требованиями НК РФ и разъяснениями ФНС. Дата проверки: 04.06.2026.",
    )
    add_callout(
        doc,
        "Вывод.",
        "Критичная ошибка расчета ставки НДС 22% в проверенной части не обнаружена: ставка и расчетная ставка 22/122 есть в перечислениях и выбираются по дате. Основной комплаенс-риск относится к регламентированной декларации НДС за периоды 2026 года: в метаданных отчета явно видны формы/настройки до 2025Кв1, тогда как ФНС ввела изменения 2026 года и новые контрольные соотношения.",
        COLORS["pale_blue"],
    )

    add_h1(doc, "1. Объем и ограничения проверки")
    add_body(doc, "Проверка выполнена как статический аудит доступной конфигурации Accounting/БухгалтерияПредприятияКОРП через MCP-сервер 1c-meta. Проверялись метаданные, регистры, макеты, перечисления и выборочные тела процедур/функций BSL.")
    add_body(doc, "Не проверялись фактические данные базы, заполненность регистров по конкретным организациям, работоспособность внешнего сервиса 1С-Отчетность, актуальность поставки регламентированной отчетности вне метаданных, а также отправка файлов в ФНС.")
    add_bullet(doc, "MCP-сервер: 1c-meta.")
    add_bullet(doc, "Ключевые запросы: list_enum_values, list_layouts, search_metadata_by_description, list_module_routines, get_routine_body.")
    add_bullet(doc, "Официальные источники: сайт ФНС России и ссылки ФНС на pravo.gov.ru/publication.pravo.gov.ru.")

    add_h1(doc, "2. Нормативная рамка на дату проверки")
    normative_rows = [
        ["Ставки НДС", "С 2026 года общая ставка - 22%; 0% и 10% применяются для специальных перечней/операций; для УСН действуют 5% и 7%; расчетные ставки включают 5/105, 7/107, 10/110 и 22/122.", "Ст. 164 НК РФ; раздел НДС на nalog.gov.ru"],
        ["Декларация НДС", "Форма, порядок и электронные форматы утверждены приказом ФНС от 05.11.2024 N ЕД-7-3/989@; изменения для налоговых периодов 2026 года утверждены приказом ФНС от 18.12.2025 N ЕД-7-3/1227@.", "ФНС: страница НДС; приказ ЕД-7-3/989@"],
        ["Контрольные соотношения 2026", "ФНС письмом от 20.02.2026 N СД-36-3/1314@ направила новые контрольные соотношения к декларации НДС 2026.", "ФНС: письмо СД-36-3/1314@"],
        ["Счета-фактуры, книги покупок/продаж", "Постановлением Правительства РФ от 23.01.2026 N 26 изменены формы и правила заполнения счета-фактуры, книги покупок и книги продаж; появилась строка 5б и новые графы 7а/11а.", "ФНС: новость от 22.04.2026"],
        ["Прослеживаемость", "По прослеживаемым товарам счета-фактуры/УПД должны содержать РНПТ, единицу и количество прослеживаемости; сведения включаются в книги и декларацию/отчетность.", "ФНС: национальная система прослеживаемости; письмо ЕА-4-15/527@"],
        ["Вычеты и восстановление", "Вычет возможен при принятии к учету и наличии первичных документов/счета-фактуры; восстановление требуется по основаниям ст. 170 НК РФ.", "Раздел НДС на nalog.gov.ru"],
    ]
    add_table(doc, ["Область", "Требование", "Источник"], normative_rows, [1800, 5400, 2160])

    add_h1(doc, "3. Что подтверждено в конфигурации")
    ok_rows = [
        ["Ставки 2026", "Перечисление СтавкиНДС содержит НДС22 и НДС22_122, а также НДС5/5_105 и НДС7/7_107.", "ОК"],
        ["Выбор ставки по дате", "Функция УчетНДСКлиентСервер.ОбщаяСтавкаНДС возвращает НДС18 до 2019, НДС20 до 01.01.2026 и НДС22 начиная с 01.01.2026.", "ОК"],
        ["Расчетная ставка по дате", "Функция ОбщаяРасчетнаяСтавкаНДС возвращает НДС22_122 начиная с 01.01.2026.", "ОК"],
        ["Книга покупок/продаж", "Есть регистры НДСЗаписиКнигиПокупок и НДСЗаписиКнигиПродаж со ставкой НДС, суммой без НДС, суммой НДС, доплистами и корректируемым периодом.", "ОК"],
        ["Документы НДС", "Найдены документы формирования записей книг покупок/продаж, подтверждения нулевой ставки, восстановления НДС, раздела 7 декларации, налогового агента.", "ОК"],
        ["Прослеживаемость", "Есть объекты УведомлениеОВвозеПрослеживаемыхТоваров, УведомлениеОбОстаткахПрослеживаемыхТоваров, ПомощникПолученияРНПТ, макеты счетов-фактур 2026 с признаком Прослеж.", "ОК"],
    ]
    add_table(doc, ["Проверка", "Доказательство из 1c-meta", "Статус"], ok_rows, [2100, 6000, 1260], risk_column=2)

    add_h1(doc, "4. Найденные риски и возможные несоответствия")
    findings = [
        [
            "НДС-01",
            "Декларация НДС 2026",
            "Высокий",
            "В списке макетов отчета РегламентированныйОтчетНДС есть ПечатныйБланк/ФормаОтчета/СоставПоказателей/Списки только до 2025Кв1; запрос find_routines_by_name по 2026 не показал процедур декларации НДС 2026.",
            "ФНС требует применять изменения за налоговые периоды 2026 года и новые КС. Возможны отказ приема, ошибки камерального контроля или искажение разделов 8-12.",
            "Добавить/подтвердить форму, состав показателей, схемы выгрузки и КС декларации НДС 2026. Если форма поставляется внешней регламентированной отчетностью, обеспечить, чтобы MCP-индексация видела актуальную поставку. Сделать автотест выгрузки декларации за 1К2026.",
        ],
        [
            "НДС-02",
            "Счет-фактура и книги 2026",
            "Средний",
            "Макеты 2026 найдены: ПФ_MXL_СчетФактура2026, ПФ_MXL_КорректировочныйСчетФактура2026, КнигаПокупок2026, КнигаПродаж2026. Но синонимы/комментарии содержат пустое место в фразе 'в редакции постановления N'.",
            "ФНС указывает новые формы 2026: строка 5б, графы 7а и 11а, отдельное отражение ставок 22/20/18. Риск не столько в наличии макета, сколько в выборе правильного макета и полноте реквизитов.",
            "Проверить алгоритм выбора макетов по дате документа/налоговому периоду. Обновить синонимы/комментарии на постановление РФ от 23.01.2026 N 26. Добавить печатные тесты для аванса, отгрузки в счет аванса, ставки 22%, прослеживаемого товара.",
        ],
        [
            "НДС-03",
            "Устаревшие ссылки в пользовательской справке",
            "Средний",
            "Справка документа ТаможеннаяДекларацияЭкспорт ссылается на приказ ФНС от 29.10.2014 N ММВ-7-3/558@. Приказ ЕД-7-3/989@ признает этот приказ утратившим силу.",
            "Пользователь может выбрать устаревший порядок/коды при подтверждении ставки 0% и заполнении реестров, даже если расчетный код актуален.",
            "Актуализировать help_text/справку по экспортным документам: приказ ЕД-7-3/989@, изменения 2026, актуальные приложения и коды операций. Добавить проверку ссылок на утратившие силу НПА в CI индекса метаданных.",
        ],
        [
            "НДС-04",
            "Налоговый агент: устаревшее описание ставки",
            "Низкий",
            "Справка обработки РегистрацияСчетовФактурНалоговогоАгента говорит, что автоматически подставляется ставка 18%. Код процедуры ЗаполнитьСтрокиДокумента фактически использует ОбщаяРасчетнаяСтавкаНДС(Дата), то есть 22/122 с 2026.",
            "Расчетный риск низкий, но подсказка вводит пользователя и аудитора в заблуждение; возможно неверное ручное изменение ставки.",
            "Исправить текст справки обработки и добавить unit-сценарий: платеж налогового агента в 2026 году формирует строку со ставкой 22/122 и корректной суммой НДС.",
        ],
        [
            "НДС-05",
            "Прослеживаемость в книгах и декларации",
            "Средний",
            "Объекты прослеживаемости и процедура выгрузки сведений о прослеживаемом товаре в РегламентированныйОтчетНДС найдены. Но по статической структуре регистров НДСЗаписиКнигиПокупок/Продаж не видно самих РНПТ-полей; вероятно, данные идут из связанных регистров/табличных частей.",
            "ФНС относит отсутствие реквизитов прослеживаемости в счетах-фактурах, книгах и декларации к типовым ошибкам. Без end-to-end теста нельзя подтвердить полноту маршрута РНПТ.",
            "Сделать сквозной тест: поступление прослеживаемого товара с РНПТ -> реализация -> счет-фактура/УПД -> книга продаж -> разделы 8/9 декларации или отчет об операциях. Проверить графы 11, 12, 12а, 13 счета-фактуры и 20-23 книги продаж.",
        ],
    ]
    add_table(doc, ["ID", "Зона", "Риск", "Доказательство", "Почему это важно", "Задание"], findings, [760, 1320, 900, 2300, 1970, 2110], risk_column=2)

    add_h1(doc, "5. Карта соответствия требованиям")
    matrix_rows = [
        ["22% и 22/122 с 01.01.2026", "Подтверждено", "Перечисления и функции выбора ставки по дате реализованы.", "Регрессионный тест на документы реализации/аванса в 2026."],
        ["Спецставки УСН 5% и 7%", "Частично подтверждено", "Ставки есть в перечислении; функциональные опции УСН с НДС найдены.", "Проверить расчет лимитов доходов и выбор режима 5/7 на данных организаций."],
        ["Счет-фактура 2026, строка 5б", "Частично подтверждено", "Макеты 2026 найдены, но выбор и заполнение строки 5б не доказаны.", "Печатный тест по авансу и отгрузке в счет аванса."],
        ["Книга покупок 2026, графа 7а", "Частично подтверждено", "КнигаПокупок2026 найдена; заполнение графы 7а не доказано.", "Тест вычета НДС с аванса."],
        ["Книга продаж 2026, графа 11а и ставки 22/20/18", "Частично подтверждено", "КнигаПродаж2026 найдена; заполнение 11а и граф по ставкам не доказано.", "Тест продажи в счет аванса и смешанных ставок."],
        ["Декларация НДС 2026", "Риск несоответствия", "В метаданных отчета видны формы до 2025Кв1; нет явных макетов/процедур 2026.", "Обновить или подтвердить внешнюю поставку регламентированной отчетности."],
        ["Прослеживаемость", "Частично подтверждено", "Есть документы РНПТ, уведомления, макеты счетов-фактур с прослеживаемостью.", "Сквозной тест маршрута РНПТ до декларации/отчета."],
        ["Раздельный учет и восстановление НДС", "Подтверждено на уровне объектов", "Найдены РаспределениеНДС, ВосстановлениеНДС, раздел 7, регистры косвенных расходов.", "Проверить бизнес-настройки организаций и учетную политику."],
    ]
    add_table(doc, ["Требование", "Статус", "Основание", "Дальнейшая проверка"], matrix_rows, [2450, 1450, 3300, 2160], risk_column=1)

    add_h1(doc, "6. Рекомендованный план исправлений")
    add_h2(doc, "Спринт 1: декларация НДС 2026")
    add_bullet(doc, "Сверить состав макетов, XML-схем, списков значений и контрольных соотношений с приказом ЕД-7-3/989@ в редакции ЕД-7-3/1227@.")
    add_bullet(doc, "Добавить автоматический сценарий формирования декларации за 1 квартал 2026: ОСНО 22%, УСН 5/7, авансы, налоговый агент, прослеживаемость.")
    add_bullet(doc, "Проверить, что MCP-индекс 1c-meta после обновления показывает 2026-специфичные макеты/настройки или однозначный механизм применения 2026 редакции.")

    add_h2(doc, "Спринт 2: счета-фактуры и книги")
    add_bullet(doc, "Проверить выбор макетов 2026 по дате документа и налоговому периоду.")
    add_bullet(doc, "В печатных формах проверить строку 5б, графу 7а книги покупок и графу 11а книги продаж.")
    add_bullet(doc, "Уточнить синонимы и комментарии макетов 2026: указать постановление РФ от 23.01.2026 N 26.")

    add_h2(doc, "Спринт 3: справка и управляемость аудита")
    add_bullet(doc, "Заменить устаревшие ссылки на приказ ММВ-7-3/558@ в справке НДС-документов.")
    add_bullet(doc, "Обновить подсказку обработки счетов-фактур налогового агента: ставка выбирается по дате, а не фиксированная 18%.")
    add_bullet(doc, "Добавить статический контроль метаданных: поиск утративших силу НПА и старых ставок в help_text/comment/synonym.")

    add_h1(doc, "7. Приложение: доказательства MCP")
    mcp_rows = [
        ["list_enum_values СтавкиНДС", "БезНДС, НДС0, НДС10, НДС10_110, НДС18, НДС18_118, НДС20, НДС20_120, НДС22, НДС22_122, НДС5, НДС5_105, НДС7, НДС7_107."],
        ["get_routine_body ОбщаяСтавкаНДС", "До 2019 возвращается НДС18; до 01.01.2026 - НДС20; с 01.01.2026 - НДС22."],
        ["get_routine_body ОбщаяРасчетнаяСтавкаНДС", "До 2019 возвращается НДС18_118; до 01.01.2026 - НДС20_120; с 01.01.2026 - НДС22_122."],
        ["list_layouts РегламентированныйОтчетНДС", "Печатные бланки, формы отчета, состав показателей, списки и настройки сравнения явно представлены до 2025Кв1; схемы выгрузки 501-511."],
        ["list_objects_by_name 2026", "Найдены ЖурналУчетаСчетовФактур2026, КнигаПокупок2026, КнигаПродаж2026, ПФ_MXL_СчетФактура2026, ПФ_MXL_СчетФактура2026Прослеж и корректировочные формы."],
        ["search_metadata_by_description РНПТ", "Найдены ПомощникПолученияРНПТ, УведомлениеОВвозеПрослеживаемыхТоваров, УведомлениеОбОстаткахПрослеживаемыхТоваров, УведомлениеОПеремещенииПрослеживаемыхТоваров."],
    ]
    add_table(doc, ["Запрос", "Результат"], mcp_rows, [2900, 6460])

    add_h1(doc, "8. Официальные источники")
    sources = [
        "ФНС России. Налог на добавленную стоимость (НДС): https://www.nalog.gov.ru/rn77/taxation/taxes/nds/",
        "ФНС России. Приказ от 05.11.2024 N ЕД-7-3/989@: https://www.nalog.gov.ru/rn77/about_fts/docs/15545040/",
        "ФНС России. Письмо от 20.02.2026 N СД-36-3/1314@ о новых КС к декларации НДС 2026: https://www.nalog.gov.ru/rn77/about_fts/about_nalog/16605831/",
        "ФНС России. Новые формы счета-фактуры, книги покупок и книги продаж в 2026 году: https://www.nalog.gov.ru/rn79/news/activities_fts/16619429/",
        "ФНС России. Национальная система прослеживаемости импортных товаров: https://www.nalog.gov.ru/rn77/related_activities/spt/",
        "ФНС России. Типичные ошибки в отчетности по прослеживаемости, письмо ЕА-4-15/527@: https://www.nalog.gov.ru/rn50/news/tax_doc_news/11830428/",
    ]
    for src in sources:
        add_bullet(doc, src)

    add_callout(
        doc,
        "Итог для владельца продукта.",
        "По НДС конфигурация выглядит зрелой в базовом учете и переходе на 22%, но для уверенного заключения о законодательном соответствии нужно закрыть доказательственный разрыв по декларации НДС 2026 и провести сквозные сценарные тесты по новым формам и прослеживаемости.",
        COLORS["green"],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
